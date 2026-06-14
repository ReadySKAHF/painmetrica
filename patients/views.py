import json
import os
from datetime import datetime, date
from io import BytesIO

from django.conf import settings
from django.contrib.auth.mixins import LoginRequiredMixin
from django.core.cache import cache
from django.core.paginator import Paginator
from django.db.models import prefetch_related_objects
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils.decorators import method_decorator
from django.views import View
from django.views.generic import DetailView
from django_ratelimit.decorators import ratelimit

from accounts.mixins import DoctorRequiredMixin
from medications.models import Medication
from patients.models import Conclusion, ConclusionMedication, ConclusionTherapy, Patient
from patients.recommendation_service import get_recommendation_context
from therapy.models import Therapy
from tests.views import _load_score_ranges, _match_score_range, _get_pathotype_label



class PatientDetailView(LoginRequiredMixin, DetailView):
    """Карточка пациента — доступна доктору (свои пациенты) и самому пациенту"""

    model = Patient
    template_name = 'patients/patient_detail.html'
    context_object_name = 'patient'

    def get_queryset(self):
        user = self.request.user
        if user.user_type == 'doctor':
            return Patient.objects.filter(
                assigned_doctor=user
            ).select_related('user', 'user__patient_profile')
        elif user.user_type == 'patient':
            return Patient.objects.filter(
                user=user
            ).select_related('user', 'user__patient_profile')
        return Patient.objects.none()

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        from collections import Counter
        from datetime import date
        from django.db.models import Prefetch
        from tests.models import Answer, ScoreRange, Test, TestSession

        DURATION_LABELS = {
            'acute':    'Острая (до 1 месяца)',
            'subacute': 'Подострая (1 - 3 месяца)',
            'chronic':  'Хроническая (более 3 месяцев)',
        }

        # Базовый QuerySet результатов (без prefetch — для счётчиков и фильтров)
        base_results_qs = (
            self.object.test_results
            .filter(status='completed')
            .select_related('test', 'session')
            .order_by('-completed_at')
        )

        # Текущий статус — первый результат с полным prefetch answers
        current_result = (
            base_results_qs
            .prefetch_related(
                Prefetch(
                    'answers',
                    queryset=Answer.objects.select_related('question__stage'),
                )
            )
            .first()
        )

        # ScoreRange для current_result
        score_ranges_for_current = {}
        if current_result:
            for sr in ScoreRange.objects.filter(test_id=current_result.test_id):
                score_ranges_for_current.setdefault(sr.test_id, []).append(sr)

        def build_sub_results(result, score_ranges_by_test):
            step_scores = {}
            step_stages = {}
            for answer in result.answers.all():
                stage = answer.question.stage
                if stage is None:
                    continue
                step = stage.sidebar_step
                if step not in step_scores:
                    step_scores[step] = 0
                    step_stages[step] = stage
                step_scores[step] += answer.score
            score_ranges = score_ranges_by_test.get(result.test_id, [])
            sub_results = []
            for step in sorted(step_scores.keys()):
                score = step_scores[step]
                stage = step_stages[step]
                score_range = _match_score_range(score_ranges, step, score)
                sub_results.append({
                    'name': stage.name,
                    'score': score,
                    'label': score_range.label if score_range else '—',
                })
            return sub_results

        if current_result:
            context['current_status'] = {
                'result': current_result,
                'sub_results': build_sub_results(current_result, score_ranges_for_current),
            }
        else:
            context['current_status'] = None

        # Фильтр по датам
        date_from_str = self.request.GET.get('date_from', '').strip()
        date_to_str = self.request.GET.get('date_to', '').strip()
        date_from = date_to = None
        if date_from_str:
            try:
                date_from = date.fromisoformat(date_from_str)
            except ValueError:
                date_from_str = ''
        if date_to_str:
            try:
                date_to = date.fromisoformat(date_to_str)
            except ValueError:
                date_to_str = ''
        context['filter_date_from'] = date_from_str
        context['filter_date_to'] = date_to_str

        # История — все результаты кроме первого (current_result), с БД-пагинацией
        history_qs = base_results_qs
        if current_result:
            history_qs = history_qs.exclude(pk=current_result.pk)
        if date_from:
            history_qs = history_qs.filter(completed_at__date__gte=date_from)
        if date_to:
            history_qs = history_qs.filter(completed_at__date__lte=date_to)

        # Наличие результатов для кнопок скачивания
        context['has_any_results'] = bool(current_result) or base_results_qs.exists()
        export_qs = base_results_qs
        if date_from:
            export_qs = export_qs.filter(completed_at__date__gte=date_from)
        if date_to:
            export_qs = export_qs.filter(completed_at__date__lte=date_to)
        context['export_has_results'] = export_qs.exists()

        # Пагинация истории на уровне БД
        history_paginator = Paginator(history_qs, 3)
        history_page_obj = history_paginator.get_page(self.request.GET.get('page', 1))
        page_results = list(history_page_obj)
        prefetch_related_objects(page_results, Prefetch(
            'answers',
            queryset=Answer.objects.select_related('question__stage'),
        ))

        # ScoreRange для страницы истории
        history_test_ids = {r.test_id for r in page_results}
        score_ranges_for_history = {}
        if history_test_ids:
            for sr in ScoreRange.objects.filter(test_id__in=history_test_ids):
                score_ranges_for_history.setdefault(sr.test_id, []).append(sr)

        context['history_page'] = history_page_obj
        context['history_results'] = [
            {'result': r, 'sub_results': build_sub_results(r, score_ranges_for_history)}
            for r in page_results
        ]

        # Проверяем возможность сравнения: ≥2 завершённых теста одной категории
        COMPARABLE = ['complex', 'painad', 'ncsr']
        category_counts = Counter(
            r['test__category']
            for r in base_results_qs.values('test__category')
            if r['test__category'] in COMPARABLE
        )
        context['can_compare'] = any(cnt >= 2 for cnt in category_counts.values())

        context['active_tab'] = self.request.GET.get('tab', 'profile')
        context['is_doctor'] = self.request.user.user_type == 'doctor'
        context['pain_duration_label'] = DURATION_LABELS.get(
            self.object.pain_duration, '—'
        ) if self.object.pain_duration else '—'
        try:
            context['date_of_birth'] = self.object.user.patient_profile.date_of_birth
        except Exception:
            context['date_of_birth'] = None

        tests = cache.get('active_tests')
        if tests is None:
            tests = list(Test.objects.filter(is_active=True).order_by('pk'))
            cache.set('active_tests', tests, 3600)
        active_sessions_map = {}
        for session in TestSession.objects.filter(
            patient=self.object,
            status='in_progress',
            taken_by=self.request.user,
        ).order_by('-started_at'):
            if session.test_id not in active_sessions_map:
                active_sessions_map[session.test_id] = session

        context['tests_with_sessions'] = [
            (t, active_sessions_map.get(t.pk)) for t in tests
        ]
        return context



class PatientWelcomeView(LoginRequiredMixin, View):
    """Приветственная страница для пациента при первом входе"""

    def get(self, request):
        user = request.user
        if user.user_type != 'patient':
            return redirect('core:dashboard')

        try:
            patient = user.patient_record
            profile = user.patient_profile
        except Exception:
            return redirect('core:dashboard')

        profile.has_seen_welcome = True
        profile.save(update_fields=['has_seen_welcome'])

        return render(request, 'patients/patient_welcome.html', {
            'patient': patient,
            'assigned_doctor': patient.assigned_doctor,
        })


class PatientHistoryView(LoginRequiredMixin, View):
    """Отдельная страница истории тестирования пациента"""

    def get(self, request, pk):
        from collections import Counter
        from datetime import date
        from django.db.models import Prefetch
        from django.http import Http404
        from tests.models import Answer, ScoreRange, Test, TestSession

        user = request.user
        if user.user_type == 'doctor':
            patient = get_object_or_404(
                Patient.objects.select_related('user', 'user__patient_profile'),
                pk=pk, assigned_doctor=user,
            )
        elif user.user_type == 'patient':
            patient = get_object_or_404(
                Patient.objects.select_related('user', 'user__patient_profile'),
                pk=pk, user=user,
            )
        else:
            raise Http404

        all_results = list(
            patient.test_results.filter(status='completed')
            .select_related('test', 'session')
            .prefetch_related(
                Prefetch('answers', queryset=Answer.objects.select_related('question__stage'))
            )
            .order_by('-completed_at')
        )

        test_ids = {r.test_id for r in all_results}
        score_ranges_by_test = {}
        if test_ids:
            for sr in ScoreRange.objects.filter(test_id__in=test_ids):
                score_ranges_by_test.setdefault(sr.test_id, []).append(sr)

        def build_sub_results(result):
            step_scores = {}
            step_stages = {}
            for answer in result.answers.all():
                stage = answer.question.stage
                if stage is None:
                    continue
                step = stage.sidebar_step
                if step not in step_scores:
                    step_scores[step] = 0
                    step_stages[step] = stage
                step_scores[step] += answer.score
            score_ranges = score_ranges_by_test.get(result.test_id, [])
            sub_results = []
            for step in sorted(step_scores.keys()):
                score = step_scores[step]
                stage = step_stages[step]
                score_range = _match_score_range(score_ranges, step, score)
                sub_results.append({
                    'name': stage.name,
                    'score': score,
                    'label': score_range.label if score_range else '—',
                })
            return sub_results

        current_result = all_results[0] if all_results else None
        if current_result:
            current_status = {'result': current_result, 'sub_results': build_sub_results(current_result)}
            history_list = all_results[1:]
        else:
            current_status = None
            history_list = all_results

        date_from_str = request.GET.get('date_from', '').strip()
        date_to_str = request.GET.get('date_to', '').strip()
        date_from = date_to = None
        if date_from_str:
            try:
                date_from = date.fromisoformat(date_from_str)
            except ValueError:
                date_from_str = ''
        if date_to_str:
            try:
                date_to = date.fromisoformat(date_to_str)
            except ValueError:
                date_to_str = ''

        history_filtered = history_list
        if date_from:
            history_filtered = [r for r in history_filtered if r.completed_at.date() >= date_from]
        if date_to:
            history_filtered = [r for r in history_filtered if r.completed_at.date() <= date_to]

        paginator = Paginator(history_filtered, 3)
        history_page = paginator.get_page(request.GET.get('page', 1))
        history_results = [
            {'result': r, 'sub_results': build_sub_results(r)}
            for r in history_page
        ]

        COMPARABLE = ['complex', 'painad', 'ncsr']
        category_counts = Counter(
            r.test.category for r in all_results if r.test.category in COMPARABLE
        )
        can_compare = any(cnt >= 2 for cnt in category_counts.values())

        return render(request, 'patients/patient_test_history.html', {
            'patient': patient,
            'current_status': current_status,
            'history_results': history_results,
            'history_page': history_page,
            'filter_date_from': date_from_str,
            'filter_date_to': date_to_str,
            'can_compare': can_compare,
            'is_doctor': user.user_type == 'doctor',
        })


class PatientUpdateAPIView(LoginRequiredMixin, View):
    """AJAX: обновить данные пациента (доктор или сам пациент)"""

    DURATION_LABELS = {
        'acute':    'Острая (до 1 месяца)',
        'subacute': 'Подострая (1 - 3 месяца)',
        'chronic':  'Хроническая (более 3 месяцев)',
    }

    def post(self, request, pk):
        user = request.user
        try:
            if user.user_type == 'doctor':
                patient = Patient.objects.select_related(
                    'user', 'user__patient_profile'
                ).get(pk=pk, assigned_doctor=user)
            elif user.user_type == 'patient':
                patient = Patient.objects.select_related(
                    'user', 'user__patient_profile'
                ).get(pk=pk, user=user)
            else:
                return JsonResponse({'success': False, 'error': 'Нет доступа'}, status=403)
        except Patient.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Пациент не найден'}, status=404)

        try:
            data = json.loads(request.body)
        except Exception:
            return JsonResponse({'success': False, 'error': 'Неверный формат данных'}, status=400)

        first_name = data.get('first_name', '')[:150].strip()
        last_name = data.get('last_name', '')[:150].strip()
        if not first_name or not last_name:
            return JsonResponse({'success': False, 'error': 'Имя и фамилия обязательны'}, status=400)

        patient_user = patient.user
        patient_user.first_name = first_name
        patient_user.middle_name = data.get('middle_name', '')[:150].strip()
        patient_user.last_name = last_name
        patient_user.save(update_fields=['first_name', 'middle_name', 'last_name'])

        # Обновляем дату рождения
        dob_str = data.get('date_of_birth', '').strip()
        try:
            profile = patient_user.patient_profile
        except Exception:
            from accounts.models import PatientProfile
            profile = PatientProfile.objects.create(user=patient_user)

        if dob_str:
            try:
                profile.date_of_birth = datetime.strptime(dob_str, '%Y-%m-%d').date()
            except ValueError:
                return JsonResponse({'success': False, 'error': 'Неверный формат даты'}, status=400)
        else:
            profile.date_of_birth = None
        profile.save(update_fields=['date_of_birth'])

        # Доктор может обновлять медицинские поля
        if user.user_type == 'doctor':
            patient.medical_history = data.get('diagnosis', '')[:5000].strip()
            patient.pain_location = data.get('pain_location', '')[:500].strip()
            duration = data.get('pain_duration', '').strip()
            valid_durations = ['', 'acute', 'subacute', 'chronic']
            patient.pain_duration = duration if duration in valid_durations else ''
            patient.save(update_fields=['medical_history', 'pain_location', 'pain_duration'])

        return JsonResponse({
            'success': True,
            'full_name': patient_user.get_full_name(),
            'last_name': patient_user.last_name,
            'first_name': patient_user.first_name,
            'date_of_birth': profile.date_of_birth.strftime('%d.%m.%Y') if profile.date_of_birth else '',
            'date_of_birth_iso': profile.date_of_birth.strftime('%Y-%m-%d') if profile.date_of_birth else '',
            'diagnosis': patient.medical_history,
            'pain_location': patient.pain_location,
            'pain_duration': patient.pain_duration,
            'pain_duration_label': (
                self.DURATION_LABELS.get(patient.pain_duration, '—')
                if patient.pain_duration else '—'
            ),
        })


class PatientMyProfileView(LoginRequiredMixin, View):
    """Редирект пациента на его собственную карточку."""

    def get(self, request):
        try:
            patient = request.user.patient_record
            return redirect('patients:detail', pk=patient.pk)
        except Exception:
            return redirect('core:dashboard')


class PatientExportExcelView(LoginRequiredMixin, View):
    """Выгрузка результатов тестов пациента в Excel.

    GET-параметры date_from / date_to (YYYY-MM-DD) фильтруют период —
    те же параметры, что используются на странице карточки пациента.
    """

    DURATION_LABELS = {
        'acute':    'Острая (до 1 месяца)',
        'subacute': 'Подострая (1 - 3 месяца)',
        'chronic':  'Хроническая (более 3 месяцев)',
    }

    def _get_patient(self, user, pk):
        """Возвращает пациента только если у пользователя есть права на просмотр."""
        if user.user_type == 'doctor':
            return get_object_or_404(
                Patient.objects.select_related(
                    'user', 'user__patient_profile', 'assigned_doctor'
                ),
                pk=pk, assigned_doctor=user,
            )
        if user.user_type == 'patient':
            return get_object_or_404(
                Patient.objects.select_related(
                    'user', 'user__patient_profile', 'assigned_doctor'
                ),
                pk=pk, user=user,
            )
        from django.http import Http404
        raise Http404

    def get(self, request, pk):
        import openpyxl
        from datetime import date as date_type
        from openpyxl.styles import Alignment

        patient = self._get_patient(request.user, pk)

        # ── Разбираем фильтр дат ──────────────────────────────────────────
        date_from_str = request.GET.get('date_from', '').strip()
        date_to_str   = request.GET.get('date_to',   '').strip()
        date_from = date_to = None

        if date_from_str:
            try:
                date_from = date_type.fromisoformat(date_from_str)
            except ValueError:
                date_from_str = ''

        if date_to_str:
            try:
                date_to = date_type.fromisoformat(date_to_str)
            except ValueError:
                date_to_str = ''

        # ── Строим queryset результатов ───────────────────────────────────
        results_qs = (
            patient.test_results
            .filter(status='completed')
            .select_related('test')
            .prefetch_related('answers__question__stage')
            .order_by('completed_at')
        )
        if date_from:
            results_qs = results_qs.filter(completed_at__date__gte=date_from)
        if date_to:
            results_qs = results_qs.filter(completed_at__date__lte=date_to)

        results = list(results_qs)

        # ── Загружаем шаблон (read-only — объект не кешируется) ───────────
        template_path = os.path.join(settings.BASE_DIR, 'test-result-template.xlsx')
        if not os.path.exists(template_path):
            raise FileNotFoundError(
                f'Файл шаблона Excel не найден: {template_path}\n'
                'Скопируйте test-result-template.xlsx в корень проекта (рядом с manage.py).'
            )
        wb = openpyxl.load_workbook(template_path)

        # ── Лист 1: Личная информация ─────────────────────────────────────
        ws_info = wb['Личная информация']

        try:
            dob = patient.user.patient_profile.date_of_birth
            dob_str = dob.strftime('%d.%m.%Y') if dob else '—'
        except Exception:
            dob_str = '—'

        doctor = patient.assigned_doctor
        ws_info['B2'] = patient.user.last_name  or '—'
        ws_info['B3'] = patient.user.first_name or '—'
        ws_info['B4'] = getattr(patient.user, 'middle_name', '') or '—'
        ws_info['B5'] = dob_str
        ws_info['B6'] = patient.medical_history or '—'
        ws_info['B7'] = patient.pain_location   or '—'
        ws_info['B8'] = (
            self.DURATION_LABELS.get(patient.pain_duration, '—')
            if patient.pain_duration else '—'
        )
        ws_info['B9'] = doctor.get_full_name() if doctor else '—'

        # ── Лист 2: Результаты тестов ─────────────────────────────────────
        ws_res = wb['Результаты тестов']

        # Переименовываем/добавляем заголовки колонок
        ws_res['C2'] = 'Тест ВАШ (VAS), NCS-R, PAINAD'
        ws_res['F2'] = 'Тест HADS. Оценка уровня тревоги'
        ws_res['G2'] = 'Тест HADS. Оценка уровня депрессии'

        # Строка 3 — пустая строка-пример из шаблона; убираем её
        ws_res.delete_rows(3)

        from openpyxl.styles import Border, Side

        center      = Alignment(horizontal='center', vertical='center', wrap_text=True)
        thin_side   = Side(style='thin')
        thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)

        # Категории, которые не имеют отдельных шагов DN4/CSI/HADS
        SINGLE_SCORE_CATEGORIES = {'ncsr', 'painad'}

        for result in results:
            # Считаем баллы по шагам сайдбара
            step_scores: dict[int, int] = {}
            for answer in result.answers.all():
                stage = answer.question.stage
                if stage is None:
                    continue
                step = stage.sidebar_step
                step_scores[step] = step_scores.get(step, 0) + answer.score

            pathotype, _ = _get_pathotype_label(
                result.test.category,
                step_scores,
                result.total_score,
                result.conclusion_label,
            )

            is_single_score = result.test.category in SINGLE_SCORE_CATEGORIES
            row = [
                result.completed_at.strftime('%d.%m.%Y %H:%M'),
                pathotype,
                step_scores.get(1) if 1 in step_scores else '—',  # VAS / NCS-R / PAINAD
                '—' if is_single_score else (step_scores.get(2) if 2 in step_scores else '—'),  # DN4
                '—' if is_single_score else (step_scores.get(3) if 3 in step_scores else '—'),  # CSI
                '—' if is_single_score else (step_scores.get(4) if 4 in step_scores else '—'),  # HADS тревога
                '—' if is_single_score else (step_scores.get(5) if 5 in step_scores else '—'),  # HADS депрессия
            ]

            row_idx = ws_res.max_row + 1
            for col_idx, value in enumerate(row, 1):
                cell = ws_res.cell(row=row_idx, column=col_idx, value=value)
                cell.alignment = center
                cell.border    = thin_border

        # ── Границы на заголовочных строках обоих листов ──────────────────
        for ws in (ws_info, ws_res):
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        cell.border = thin_border

        # ── Авторазмер колонок (по длиннейшему значению) ──────────────────
        for ws in (ws_info, ws_res):
            for col_cells in ws.columns:
                max_len = 0
                col_letter = col_cells[0].column_letter
                for cell in col_cells:
                    if cell.value:
                        max_len = max(max_len, len(str(cell.value)))
                ws.column_dimensions[col_letter].width = max_len + 4

        # ── Формируем имя файла: просто ФИО пациента ─────────────────────
        safe_name = patient.user.get_full_name().replace(' ', '_')
        filename  = f'{safe_name}.xlsx'

        # ── Отдаём файл ───────────────────────────────────────────────────
        output = BytesIO()
        wb.save(output)
        output.seek(0)

        response = HttpResponse(
            output.read(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        from urllib.parse import quote
        encoded = quote(filename, safe='')
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded}"
        return response


class ConclusionListView(DoctorRequiredMixin, View):
    """История заключений пациента"""

    def get(self, request, pk):
        from datetime import datetime
        patient = get_object_or_404(
            Patient.objects.select_related('user'),
            pk=pk,
            assigned_doctor=request.user,
        )

        def parse_date(s):
            """Парсит ГГГГ-ММ-ДД (ISO), возвращает date или None"""
            if not s:
                return None
            try:
                return date.fromisoformat(s.strip())
            except ValueError:
                return None

        date_from_str = request.GET.get('date_from', '').strip()
        date_to_str = request.GET.get('date_to', '').strip()
        date_from = parse_date(date_from_str)
        date_to = parse_date(date_to_str)
        if date_from_str and not date_from:
            date_from_str = ''
        if date_to_str and not date_to:
            date_to_str = ''

        conclusions_qs = Conclusion.objects.filter(patient=patient).prefetch_related(
            'conclusion_medications__medication',
            'conclusion_therapies__therapy',
        ).order_by('-created_at')
        if date_from:
            conclusions_qs = conclusions_qs.filter(created_at__date__gte=date_from)
        if date_to:
            conclusions_qs = conclusions_qs.filter(created_at__date__lte=date_to)

        paginator = Paginator(conclusions_qs, 10)
        page = paginator.get_page(request.GET.get('page', 1))

        return render(request, 'patients/conclusion_history.html', {
            'patient': patient,
            'page': page,
            'total_count': paginator.count,
            'filter_date_from': date_from_str,
            'filter_date_to': date_to_str,
        })


class ConclusionMedicineView(DoctorRequiredMixin, View):
    """Шаг 1: выбор лекарств для заключения"""

    def _get_patient(self, request, pk):
        return get_object_or_404(
            Patient.objects.select_related('user'),
            pk=pk,
            assigned_doctor=request.user,
        )

    def _get_last_result(self, patient):
        from tests.models import TestResult
        return (
            TestResult.objects.filter(patient=patient, status='completed')
            .order_by('-completed_at')
            .first()
        )

    def get(self, request, pk):
        patient = self._get_patient(request, pk)
        last_result = self._get_last_result(patient)
        rec_ctx = get_recommendation_context(patient, last_result)

        if not rec_ctx.show_meds_page:
            if not rec_ctx.show_rehab_page:
                # Пропускаем обе страницы — создаём пустое заключение сразу
                conclusion = Conclusion.objects.create(
                    patient=patient,
                    doctor=request.user,
                    last_test_result=last_result,
                )
                return redirect(
                    reverse('patients:detail', kwargs={'pk': pk})
                    + f'?download_conclusion={conclusion.pk}'
                )
            return redirect('patients:conclusion_rehabilitation', pk=pk)

        med_types_filter = [t.strip() for t in request.GET.getlist('med_type') if t.strip()]

        # Показываем только рекомендованные лекарства
        rec_qs = Medication.objects.filter(
            pk__in=rec_ctx.recommended_med_ids,
        ).order_by('medication_type', 'name')

        medications = rec_qs.filter(medication_type__in=med_types_filter) if med_types_filter else rec_qs

        med_types = list(
            rec_qs.values_list('medication_type', flat=True)
            .distinct().exclude(medication_type='').order_by('medication_type')
        )

        # Предварительно выбранные: из сессии (при возврате назад), иначе пусто
        selected_ids = request.session.get(f'conclusion_med_ids_{pk}', [])

        return render(request, 'patients/conclusion_medicine.html', {
            'patient': patient,
            'medications': medications,
            'all_medications': rec_qs,
            'med_types': med_types,
            'selected_ids': selected_ids,
            'current_med_types': med_types_filter,
            'last_result': last_result,
            'show_rehab_page': rec_ctx.show_rehab_page,
        })

    def post(self, request, pk):
        patient = self._get_patient(request, pk)
        last_result = self._get_last_result(patient)
        rec_ctx = get_recommendation_context(patient, last_result)

        selected_ids = request.POST.getlist('medication_ids')
        med_ids = [int(i) for i in selected_ids if i.isdigit()]
        request.session[f'conclusion_med_ids_{pk}'] = med_ids

        med_schemes = {}
        for med_id in med_ids:
            scheme = request.POST.get(f'scheme_{med_id}', '').strip()
            med_schemes[str(med_id)] = scheme
        request.session[f'conclusion_med_schemes_{pk}'] = med_schemes

        if not rec_ctx.show_rehab_page:
            # Реабилитация не нужна — сразу создаём заключение
            conclusion = self._create_conclusion(request, patient, med_ids, med_schemes, last_result)
            request.session.pop(f'conclusion_med_ids_{pk}', None)
            request.session.pop(f'conclusion_med_schemes_{pk}', None)
            return redirect(
                reverse('patients:detail', kwargs={'pk': pk})
                + f'?download_conclusion={conclusion.pk}'
            )

        return redirect('patients:conclusion_rehabilitation', pk=pk)

    @staticmethod
    def _create_conclusion(request, patient, med_ids, med_schemes, last_result):
        conclusion = Conclusion.objects.create(
            patient=patient,
            doctor=request.user,
            last_test_result=last_result,
        )
        if med_ids:
            meds = Medication.objects.filter(pk__in=med_ids)
            for med in meds:
                ConclusionMedication.objects.create(
                    conclusion=conclusion,
                    medication=med,
                    custom_scheme=med_schemes.get(str(med.pk), ''),
                )
        return conclusion


class ConclusionRehabilitationView(DoctorRequiredMixin, View):
    """Шаг 2: выбор реабилитационных мероприятий и создание заключения"""

    def _get_patient(self, request, pk):
        return get_object_or_404(
            Patient.objects.select_related('user'),
            pk=pk,
            assigned_doctor=request.user,
        )

    def _get_last_result(self, patient):
        from tests.models import TestResult
        return (
            TestResult.objects.filter(patient=patient, status='completed')
            .order_by('-completed_at')
            .first()
        )

    def get(self, request, pk):
        patient = self._get_patient(request, pk)
        last_result = self._get_last_result(patient)
        rec_ctx = get_recommendation_context(patient, last_result)

        # Показываем только рекомендованные мероприятия
        therapies = Therapy.objects.filter(
            pk__in=rec_ctx.recommended_therapy_ids,
        ).order_by('therapy_type', 'name')

        # Предварительно выбранные: из сессии (при возврате назад), иначе пусто
        selected_ids = request.session.get(f'conclusion_therapy_ids_{pk}', [])

        return render(request, 'patients/conclusion_rehabilitation.html', {
            'patient': patient,
            'therapies': therapies,
            'selected_ids': selected_ids,
            'last_result': last_result,
        })

    def post(self, request, pk):
        from tests.models import TestResult

        patient = self._get_patient(request, pk)

        # Лекарства и их схемы из сессии
        med_ids = request.session.get(f'conclusion_med_ids_{pk}', [])
        med_schemes = request.session.get(f'conclusion_med_schemes_{pk}', {})

        therapy_ids = [int(i) for i in request.POST.getlist('therapy_ids') if i.isdigit()]

        # Сохраняем therapy_ids в сессию на случай возврата
        request.session[f'conclusion_therapy_ids_{pk}'] = therapy_ids

        # Последний завершённый тест пациента
        last_result = (
            TestResult.objects.filter(patient=patient, status='completed')
            .order_by('-completed_at')
            .first()
        )

        # Создаём заключение
        conclusion = Conclusion.objects.create(
            patient=patient,
            doctor=request.user,
            last_test_result=last_result,
        )

        # Добавляем лекарства с отредактированными схемами
        if med_ids:
            meds = Medication.objects.filter(pk__in=med_ids)
            for med in meds:
                custom = med_schemes.get(str(med.pk), '')
                ConclusionMedication.objects.create(
                    conclusion=conclusion,
                    medication=med,
                    custom_scheme=custom,
                )

        # Добавляем терапии с отредактированными схемами
        if therapy_ids:
            therapies_qs = Therapy.objects.filter(pk__in=therapy_ids)
            for therapy in therapies_qs:
                custom = request.POST.get(f'scheme_{therapy.pk}', '').strip()
                ConclusionTherapy.objects.create(
                    conclusion=conclusion,
                    therapy=therapy,
                    custom_scheme=custom,
                )

        # Очищаем сессию
        request.session.pop(f'conclusion_med_ids_{pk}', None)
        request.session.pop(f'conclusion_med_schemes_{pk}', None)
        request.session.pop(f'conclusion_therapy_ids_{pk}', None)

        return redirect(
            reverse('patients:detail', kwargs={'pk': pk})
            + f'?download_conclusion={conclusion.pk}'
        )


@method_decorator(ratelimit(key='user', rate='10/h', method='GET', block=True), name='get')
class ConclusionDownloadView(DoctorRequiredMixin, View):
    """Скачать заключение в формате DOCX по шаблону conclusion-template.docx"""

    DURATION_LABELS = {
        'acute':    'Острая (до 1 месяца)',
        'subacute': 'Подострая (1–3 месяца)',
        'chronic':  'Хроническая (более 3 месяцев)',
    }

    def get(self, request, pk, conclusion_pk):
        import docx
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from tests.models import ScoreRange

        patient = get_object_or_404(
            Patient.objects.select_related('user', 'user__patient_profile'),
            pk=pk,
            assigned_doctor=request.user,
        )
        conclusion = get_object_or_404(
            Conclusion.objects.select_related(
                'doctor',
                'doctor__doctor_profile',
                'last_test_result__test',
            ).prefetch_related(
                'conclusion_medications__medication',
                'conclusion_therapies__therapy',
                'last_test_result__answers__question__stage',
            ),
            pk=conclusion_pk,
            patient=patient,
        )

        # ── Результаты тестирования из last_test_result ────────────────────
        # Показываем только шаги, для которых есть совпадение в ScoreRange
        test_results_data = []
        step_scores_for_conclusion = {}
        step_labels_for_conclusion = {}
        test_category = ''
        if conclusion.last_test_result:
            result = conclusion.last_test_result
            test_category = result.test.category or ''
            score_ranges = list(ScoreRange.objects.filter(test=result.test))
            step_scores = {}
            step_stages = {}
            for answer in result.answers.all():
                stage = answer.question.stage
                if stage is None:
                    continue
                step = stage.sidebar_step
                if step not in step_scores:
                    step_scores[step] = 0
                    step_stages[step] = stage
                step_scores[step] += answer.score
            for step in sorted(step_scores.keys()):
                score = step_scores[step]
                stage = step_stages[step]
                score_range = _match_score_range(score_ranges, step, score)
                if score_range is None:
                    continue  # шаги без ScoreRange не выводим
                test_results_data.append({
                    'name': stage.name,
                    'label': score_range.label,
                    'score': score,
                })
                step_labels_for_conclusion[step] = score_range.label
            step_scores_for_conclusion = step_scores

        # ── Данные пациента ────────────────────────────────────────────────
        patient_user = patient.user
        try:
            dob = patient_user.patient_profile.date_of_birth
            dob_str = dob.strftime('%d.%m.%Y') if dob else ''
        except Exception:
            dob_str = ''

        fio_parts = filter(None, [
            patient_user.last_name,
            patient_user.first_name,
            getattr(patient_user, 'middle_name', ''),
        ])
        patient_fio = ' '.join(fio_parts)
        patient_name_line = patient_fio
        if dob_str:
            patient_name_line += f', {dob_str}'

        diagnosis = patient.medical_history or '–'
        created_date = conclusion.created_at.strftime('%d.%m.%Y')
        duration_text = self.DURATION_LABELS.get(patient.pain_duration, patient.pain_duration or '–')

        # ── Заключение ────────────────────────────────────────────────────
        from patients.conclusion_service import build_conclusion_text
        conclusion_text, show_recommendations = build_conclusion_text(
            category=test_category,
            step_scores=step_scores_for_conclusion,
            step_labels=step_labels_for_conclusion,
            pain_duration=patient.pain_duration or '',
        )

        # ── Врач: Фамилия И.О. ─────────────────────────────────────────────
        doctor = conclusion.doctor
        doc_last = doctor.last_name or ''
        doc_first = (doctor.first_name or '')[:1]
        doc_middle = (getattr(doctor, 'middle_name', '') or '')[:1]
        doctor_short = doc_last
        if doc_first:
            doctor_short += f' {doc_first}.'
        if doc_middle:
            doctor_short += f' {doc_middle}.'

        # ── Лекарства и мероприятия ────────────────────────────────────────
        medications = [
            (cm.medication.name, cm.custom_scheme or cm.medication.prescription_scheme)
            for cm in conclusion.conclusion_medications.all()
        ]
        therapies = [
            (ct.therapy.name, ct.custom_scheme or ct.therapy.prescription_scheme)
            for ct in conclusion.conclusion_therapies.all()
        ]

        # ── Данные врача для шапки ─────────────────────────────────────────
        doctor_profile = getattr(doctor, 'doctor_profile', None)
        doctor_workplace = (doctor_profile.workplace if doctor_profile else '') or ''
        doctor_position = (doctor_profile.position if doctor_profile else '') or ''
        doctor_fio_parts = filter(None, [
            doctor.last_name,
            doctor.first_name,
            getattr(doctor, 'middle_name', ''),
        ])
        doctor_fio = ' '.join(doctor_fio_parts)

        # ── Открываем шаблон ───────────────────────────────────────────────
        template_path = os.path.join(settings.BASE_DIR, 'conclusion-template.docx')
        document = docx.Document(template_path)

        # Заменяем шапку шаблона данными врача из БД
        def _replace_para_text(para, new_text):
            """Заменяет текст параграфа, сохраняя форматирование первого run."""
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)

        header_map = {
            0: doctor_workplace,
            1: doctor_position,
            2: doctor_fio,
        }
        for idx, new_text in header_map.items():
            if idx < len(document.paragraphs) and new_text:
                _replace_para_text(document.paragraphs[idx], new_text)

        # Находим параграф "КОНСУЛЬТАТИВНОЕ ЗАКЛЮЧЕНИЕ" и удаляем всё после него
        header_end = None
        for i, p in enumerate(document.paragraphs):
            if 'КОНСУЛЬТАТИВНОЕ ЗАКЛЮЧЕНИЕ' in p.text:
                header_end = i
                break

        if header_end is not None:
            for p in list(document.paragraphs[header_end + 1:]):
                p._element.getparent().remove(p._element)

        # Одинарный интервал для параграфов шапки
        for p in document.paragraphs:
            fmt = p.paragraph_format
            fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)

        # ── Вспомогательная: добавить параграф с одним или несколькими runs ─
        def _add(runs, bold=False, italic=False,
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12):
            """runs — строка ИЛИ список кортежей (text, bold, italic)."""
            para = document.add_paragraph()
            para.alignment = align
            fmt = para.paragraph_format
            fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            if isinstance(runs, str):
                if runs:
                    r = para.add_run(runs)
                    r.bold = bold
                    r.italic = italic
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(size)
            else:
                for text, r_bold, r_italic in runs:
                    if text:
                        r = para.add_run(text)
                        r.bold = r_bold
                        r.italic = r_italic
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(size)
            return para

        # ── Формируем содержимое ───────────────────────────────────────────
        _add('')

        # Дата осмотра — вся строка курсив
        _add(f'Дата осмотра: {created_date} г.',
             italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)

        # Ф.И.О. пациента — вся строка курсив
        fio_line = (f'Ф.И.О. пациента: {patient_name_line} г.'
                    if dob_str else f'Ф.И.О. пациента: {patient_name_line}')
        _add(fio_line, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)

        # Диагноз: — метка курсив, значение обычное
        _add([
            ('Диагноз: ', False, True),
            (diagnosis,   False, False),
        ])

        _add('')
        _add('')

        def _conclusion_test_name(name):
            """'Тест HADS. Оценка уровня тревоги' → 'HADS (оценка уровня тревоги)'"""
            if name.startswith('Тест '):
                name = name[5:]
            if '. ' in name:
                prefix, suffix = name.split('. ', 1)
                name = f'{prefix} ({suffix.lower()})'
            return name

        # Результаты тестирования — курсив
        _add('Результаты тестирования:', italic=True)
        for r in test_results_data:
            _add(f"Опросник {_conclusion_test_name(r['name'])}: {r['label']} ({r['score']} баллов)")

        # Длительность болевого синдрома
        _add(f'Длительность болевого синдрома: {duration_text}')

        _add('')

        _add(conclusion_text, bold=True)

        _add('')

        if show_recommendations:
            _add('Рекомендации по лечению и реабилитации:', bold=True)
            for name, scheme in medications:
                _add(f'- {name}: {scheme}', bold=True)
            for name, scheme in therapies:
                _add(f'- {name}: {scheme}', bold=True)
            _add('')

        _add(f'Врач:                      {doctor_short}', bold=True,
             align=WD_ALIGN_PARAGRAPH.RIGHT)

        # ── Отдаём файл ───────────────────────────────────────────────────
        output = BytesIO()
        document.save(output)
        output.seek(0)

        safe_name = patient_user.get_full_name().replace(' ', '_')
        filename = f'Заключение_{safe_name}_{created_date}.docx'

        response = HttpResponse(
            output.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        from urllib.parse import quote
        encoded = quote(filename, safe='')
        response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded}"
        return response
